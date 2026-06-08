"""Extract text and inline images from DOCX files to create image-caption datasets."""
from __future__ import annotations

import argparse
import json
import logging
import re
from io import BytesIO
from pathlib import Path
from typing import Iterable

from docx import Document
from PIL import Image

from vlm_weather.common import convert_doc_to_docx, ensure_dir, file_stem, find_files

logger = logging.getLogger(__name__)

DEFAULT_FIGURE_LABELS = [f"图{i}" for i in range(1, 9)]
CAPTION_TEMPLATE: dict = {"captions": [{"role": "caption", "content": ""}]}
SUPPORTED_IMAGE_EXTS = {"jpg", "jpeg", "png", "gif", "bmp", "tif"}


def extract_text(docx_path: str | Path) -> str:
    """Extract all paragraph text from a DOCX file."""
    doc = Document(str(docx_path))
    return "\n".join(para.text for para in doc.paragraphs)


def extract_paragraphs_by_figure_labels(
    text: str, figure_labels: Iterable[str]
) -> dict[str, list[str]]:
    """Find paragraphs containing ``见图N`` and group by figure label.

    Returns:
        Mapping from label (e.g., ``"图1"``) to list of cleaned paragraphs.
    """
    paragraphs = text.splitlines()
    result: dict[str, list[str]] = {label: [] for label in figure_labels}

    for label in figure_labels:
        marker = "见" + label
        pattern = re.compile(re.escape(marker), re.DOTALL)
        for para in paragraphs:
            if pattern.search(para):
                cleaned = pattern.sub("", para)
                result[label].append(cleaned)
    return result


def extract_inline_images(
    docx_path: str | Path,
    image_dir: str | Path,
    backup_dir: str | Path | None = None,
    min_height: int = 4,
) -> int:
    """Extract inline images from a DOCX file.

    Returns:
        Number of images saved.
    """
    ensure_dir(image_dir)
    if backup_dir is not None:
        ensure_dir(backup_dir)

    doc = Document(str(docx_path))
    stem = file_stem(docx_path)
    saved = 0

    for idx, shape in enumerate(doc.inline_shapes, start=1):
        content_id = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
        related_part = doc.part.related_parts[content_id]
        if not related_part.content_type.startswith("image"):
            continue

        ext = related_part.content_type.split("/")[-1].lower()
        if ext not in SUPPORTED_IMAGE_EXTS:
            logger.warning("Unsupported image type: %s", ext)
            continue

        blob = related_part._blob
        try:
            image = Image.open(BytesIO(blob))
        except Exception as exc:
            logger.error("Cannot decode image %s-图%d: %s", stem, idx, exc)
            continue

        if image.height < min_height:
            continue

        name = f"{stem}-图{idx}.{ext}"
        primary = Path(image_dir) / name
        primary.write_bytes(blob)
        saved += 1

        if backup_dir is not None:
            (Path(backup_dir) / name).write_bytes(blob)

    return saved


def write_caption_jsons(
    docx_path: str | Path,
    paragraphs: dict[str, list[str]],
    labels_dir: str | Path,
    backup_dir: str | Path | None = None,
) -> int:
    """Write caption JSON files grouped by figure label.

    Returns:
        Number of JSON files written.
    """
    ensure_dir(labels_dir)
    if backup_dir is not None:
        ensure_dir(backup_dir)

    stem = file_stem(docx_path)
    written = 0
    for label, paras in paragraphs.items():
        if not paras:
            continue
        data = json.loads(json.dumps(CAPTION_TEMPLATE))
        data["captions"][0]["content"] = paras[0]

        filename = f"{stem}-{label}.json"
        (Path(labels_dir) / filename).write_text(
            json.dumps(data, ensure_ascii=False, indent=4),
            encoding="utf-8",
        )
        written += 1

        if backup_dir is not None:
            (Path(backup_dir) / filename).write_text(
                json.dumps(data, ensure_ascii=False, indent=4),
                encoding="utf-8",
            )
    return written


def rename_doc_files_sequentially(
    file_paths: list[str | Path], start_id: str
) -> list[Path]:
    """Rename .doc/.docx files to sequential IDs.

    .doc files are first converted to .docx (Windows only).

    Returns:
        List of final .docx paths.
    """
    width = len(start_id)
    current = int(start_id)
    output_paths: list[Path] = []

    for source in file_paths:
        source = Path(source)
        new_name = f"{current:0>{width}}.docx"
        new_path = source.with_name(new_name)
        ext = source.suffix.lower()

        if ext == ".doc":
            if convert_doc_to_docx(source, new_path):
                try:
                    source.unlink()
                except OSError as exc:
                    logger.warning("Cannot remove original .doc %s: %s", source, exc)
                output_paths.append(new_path)
            else:
                continue
        else:
            if source != new_path:
                source.rename(new_path)
            output_paths.append(new_path)

        current += 1
    return output_paths


def process_directory(
    input_dir: str | Path,
    output_dir: str | Path,
    start_id: str = "800002069",
    figure_labels: Iterable[str] | None = None,
    rename_files: bool = True,
) -> None:
    """Full pipeline: convert -> rename -> extract text and images."""
    figure_labels = list(figure_labels or DEFAULT_FIGURE_LABELS)
    images_dir = Path(output_dir) / "images"
    labels_dir = Path(output_dir) / "labels"
    backup_dir = Path(output_dir) / "Img_lab"

    doc_files = find_files(input_dir, [".doc", ".docx"])
    logger.info("Found %d files to process", len(doc_files))

    if rename_files:
        docx_files = rename_doc_files_sequentially(doc_files, start_id)
    else:
        docx_files = [p for p in doc_files if p.suffix.lower() == ".docx"]

    logger.info("Processing %d .docx files", len(docx_files))
    for docx_path in docx_files:
        logger.info("--- %s", docx_path)
        text = extract_text(docx_path)
        paragraphs = extract_paragraphs_by_figure_labels(text, figure_labels)
        write_caption_jsons(docx_path, paragraphs, labels_dir, backup_dir)
        extract_inline_images(docx_path, images_dir, backup_dir)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Extract text and images from DOCX to create image-caption datasets."
    )
    parser.add_argument("--input_dir", required=True, help="Root directory with .doc/.docx files")
    parser.add_argument("--output_dir", required=True, help="Output root directory")
    parser.add_argument("--start_id", default="800002069", help="Sequential rename start ID")
    parser.add_argument("--figure_labels", nargs="*", default=None, help="Custom figure labels")
    parser.add_argument("--no_rename", action="store_true", help="Skip rename/conversion")
    args = parser.parse_args()

    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

    process_directory(
        input_dir=args.input_dir,
        output_dir=args.output_dir,
        start_id=args.start_id,
        figure_labels=args.figure_labels,
        rename_files=not args.no_rename,
    )


if __name__ == "__main__":
    main()
